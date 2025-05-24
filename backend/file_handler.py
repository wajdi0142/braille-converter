import os
import pdfplumber
from docx import Document
from reportlab.lib.units import mm
from PyQt5.QtGui import QPainter, QFont
from PyQt5.QtPrintSupport import QPrinter
from PIL import Image, ImageEnhance
import pytesseract
import numpy as np
import cv2
import re
import unicodedata
import logging

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from xml.sax.saxutils import escape
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK

logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Configuration des polices
BRAILLE_FONT_NAME = "Noto Sans Braille"
FALLBACK_FONT = "Times New Roman"
FONT_PATH = os.getenv("FONT_PATH", r"C:\Users\LENOVO\Downloads\Noto_Sans_Symbols_2\NotoSansSymbols2-Regular.ttf")

# Constants for export_pdf
TEXT_FONT_NAME = "Times New Roman"
FONT_PATHS = {
    "Times New Roman": "C:/Windows/Fonts/times.ttf",
    "Times New Roman-Bold": "C:/Windows/Fonts/timesbd.ttf",
    "Times New Roman-Italic": "C:/Windows/Fonts/timesi.ttf",
    "Noto Sans Braille": FONT_PATH,
}
DEFAULT_LINES_PER_PAGE = 25
DEFAULT_LINE_WIDTH = 120
DEFAULT_INDENT = 0
DEFAULT_LINE_SPACING = 1.0

class FileHandler:
    def __init__(self):
        self.last_gcode = None
        self.parent = None

    def extract_text(self, file_path, max_pages=10):
        try:
            if not os.path.exists(file_path):
                print(f"Erreur : Le fichier {file_path} n'existe pas.")
                return ""
            
            if file_path.lower().endswith('.txt') or file_path.lower().endswith('.bfr'):
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read()
                except Exception as e:
                    logging.error(f"Erreur lors de la lecture du fichier texte {file_path}: {str(e)}")
                    print(f"Erreur lors de la lecture du fichier texte : {os.path.basename(file_path)}")
                    return ""

            if file_path.lower().endswith('.pdf'):
                try:
                    with pdfplumber.open(file_path) as pdf:
                        text = ''
                        for i, page in enumerate(pdf.pages):
                            if i >= max_pages:
                                logging.warning(f"Limite de {max_pages} pages atteinte pour {file_path}")
                                break
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + '\n'
                        logging.debug(f"Extraction PDF réussie pour {file_path} ({len(pdf.pages)} pages)")
                        return text.strip()
                except Exception as e:
                    logging.error(f"Erreur lors de l'extraction du PDF {file_path}: {str(e)}")
                    print(f"Erreur lors de l'extraction du PDF : {os.path.basename(file_path)}")
                    return ""

            if file_path.lower().endswith('.docx'):
                try:
                    doc = Document(file_path)
                    text = '\n'.join(para.text for para in doc.paragraphs if para.text.strip())
                    logging.debug(f"Extraction DOCX réussie pour {file_path}")
                    return text
                except Exception as e:
                    logging.error(f"Erreur lors de l'extraction du DOCX {file_path}: {str(e)}")
                    print(f"Erreur lors de l'extraction du DOCX : {os.path.basename(file_path)}")
                    return ""

            print(f"Format non pris en charge : {file_path}")
            logging.warning(f"Format non pris en charge pour {file_path}")
            return ""

        except Exception as e:
            logging.critical(f"Erreur inattendue lors de l'extraction de {file_path}: {str(e)}", exc_info=True)
            print(f"Erreur inattendue lors de l'extraction de : {os.path.basename(file_path)}")
            return ""

    def _is_text_arabic(self, text):
        arabic_chars = 0
        total_chars = 0
        for char in text:
            if unicodedata.bidirectional(char) in ('AL', 'R', 'BN'):
                arabic_chars += 1
            if char.isalpha():
                total_chars += 1
        return total_chars > 0 and (arabic_chars / total_chars) > 0.5

    def _text_to_braille_french(self, text):
        braille_table = {
            'a': '⠁', 'b': '⠃', 'c': '⠉', 'd': '⠙', 'e': '⠑', 'f': '⠋', 'g': '⠛', 'h': '⠓', 'i': '⠊',
            'j': '⠚', 'k': '⠅', 'l': '⠇', 'm': '⠍', 'n': '⠝', 'o': '⠕', 'p': '⠏', 'q': '⠟', 'r': '⠗',
            's': '⠎', 't': '⠞', 'u': '⠥', 'v': '⠧', 'w': '⠺', 'x': '⠭', 'y': '⠽', 'z': '⠵',
            'é': '⠿', 'è': '⠾', 'ê': '⠻', 'ë': '⠷', 'à': '⠈⠁', 'â': '⠈⠁', 'ô': '⠈⠕', 'î': '⠈⠊',
            'ù': '⠈⠥', 'ç': '⠯', 'œ': '⠪', '-': '⠤', '/': '⠌', ' ': ' ',
            '1': '⠼⠁', '2': '⠼⠃', '3': '⠼⠉', '4': '⠼⠙', '5': '⠼⠑',
            '6': '⠼⠋', '7': '⠼⠛', '8': '⠼⠓', '9': '⠼⠊', '0': '⠼⠕',
            '.': '⠲', ',': '⠂', ';': '⠆', ':': '⠒', '!': '⠖', '?': '⠦',
            '(': '⠦', ')': '⠴', '[': '⠦', ']': '⠴', '*': '⠔', '"': '⠦',
        }
        braille_text = ''
        for char in text.lower():
            braille_text += braille_table.get(char, ' ')
        return braille_text

    def _text_to_braille_arabic(self, text):
        braille_table = {
            'ا': '⠁', 'ب': '⠃', 'ت': '⠞', 'ث': '⠹', 'ج': '⠚', 'ح': '⠓', 'خ': '⠱', 'د': '⠙', 'ذ': '⠮',
            'ر': '⠗', 'ز': '⠵', 'س': '⠎', 'ش': '⠩', 'ص': '⠯', 'ض': '⠟', 'ط': '⠷', 'ظ': '⠧', 'ع': '⠫',
            'غ': '⠻', 'ف': '⠋', 'ق': '⠭', 'ك': '⠅', 'ل': '⠇', 'م': '⠍', 'ن': '⠝', 'ه': '⠗', 'و': '⠺',
            'ي': '⠽', 'ة': '⠢', ' ': ' ',
            '1': '⠼⠁', '2': '⠼⠃', '3': '⠼⠉', '4': '⠼⠙', '5': '⠼⠑',
            '6': '⠼⠋', '7': '⠼⠛', '8': '⠼⠓', '9': '⠼⠊', '0': '⠼⠕',
            '.': '⠲', ',': '⠂', ';': '⠆', ':': '⠒', '!': '⠖', '?': '⠦',
        }
        braille_text = ''
        text = text[::-1]
        for char in text:
            braille_text += braille_table.get(char, ' ')
        return braille_text

    def convert_to_braille(self, text):
        if not text.strip():
            return "⠁⠥⠉⠥⠝ ⠞⠑⠭⠞⠑ ⠁ ⠉⠕⠝⠧⠑⠗⠞⠊⠗⠲"

        is_arabic = self._is_text_arabic(text)
        if is_arabic:
            return self._text_to_braille_arabic(text)
        else:
            return self._text_to_braille_french(text)

    def image_to_braille(self, file_path, mode='text', width=40, height=20, contrast=2.0, lang='fra+ara', psm=6):
        try:
            with open(file_path, 'rb') as f:
                image_data = np.frombuffer(f.read(), np.uint8)
            image = cv2.imdecode(image_data, cv2.IMREAD_GRAYSCALE)
            if image is None:
                raise ValueError("Impossible de charger l'image.")

            extracted_text = ""
            braille_text = ""
            graphic_braille = ""

            image_pil = Image.fromarray(image)
            enhancer = ImageEnhance.Contrast(image_pil)
            image_pil = enhancer.enhance(contrast)
            image = np.array(image_pil)

            if mode in ['text', 'hybrid']:
                try:
                    scale_factor = 3
                    image_ocr = cv2.resize(image, None, fx=scale_factor, fy=scale_factor, interpolation=cv2.INTER_CUBIC)
                    image_ocr = cv2.GaussianBlur(image_ocr, (5, 5), 0)
                    _, image_ocr = cv2.threshold(image_ocr, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                    image_ocr_pil = Image.fromarray(image_ocr)

                    extracted_text = pytesseract.image_to_string(image_ocr_pil, lang=lang, config=f'--psm {psm} --oem 1')
                    print("Texte extrait par OCR :", extracted_text)

                    if extracted_text.strip():
                        extracted_text = extracted_text.replace('•', '- COOKIES')
                        extracted_text = re.sub(r'\n+', '\n', extracted_text).strip()
                        extracted_text = re.sub(r'[!*]+', ' ', extracted_text)
                        extracted_text = extracted_text.replace('ä', 'a').replace('@', '').replace('>', '')
                        extracted_text = extracted_text.replace('0', 'o').replace('1', 'i')
                        braille_text = self.convert_to_braille(extracted_text)
                        if not braille_text.strip():
                            braille_text = "Erreur lors de la conversion du texte en Braille."
                        if mode == "text":
                            return extracted_text, braille_text
                    else:
                        print("Aucun texte extrait par l'OCR.")
                        extracted_text = "Aucun texte extrait."
                        braille_text = "⠁⠥⠉⠥⠝ ⠞⠑⠭⠞⠑ ⠑⠭⠞⠗⠁⠊⠞⠲"
                except Exception as e:
                    print(f"Erreur OCR : {e}")
                    extracted_text = "Erreur lors de l'extraction du texte."
                    braille_text = "⠑⠗⠗⠑⠥⠗ ⠇⠕⠗⠎ ⠙⠑ ⠇⠦⠑⠭⠞⠗⠁⠉⠞⠊⠕⠝⠲"

            if mode in ['graphic', 'hybrid']:
                image_graphic = cv2.resize(image, (width * 2, height * 4), interpolation=cv2.INTER_AREA)
                _, image_graphic = cv2.threshold(image_graphic, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                graphic_braille = self._image_to_braille_graphic(image_graphic, width, height)
                self.last_gcode = graphic_braille

                if mode == 'graphic':
                    return extracted_text, graphic_braille
                else:
                    combined_braille = f"{braille_text}\n\n=== Représentation graphique ===\n{graphic_braille}"
                    return extracted_text, combined_braille

            return extracted_text, braille_text

        except Exception as e:
            print(f"Erreur lors de la conversion de l'image : {e}")
            return "Erreur lors de la conversion.", "⠑⠗⠗⠑⠥⠗ ⠙⠑ ⠉⠕⠝⠧⠑⠗⠎⠊⠕⠝⠲"

    def _image_to_braille_graphic(self, image, width, height):
        braille_chars = ['⠀', '⠁', '⠂', '⠃', '⠄', '⠅', '⠆', '⠇', '⠈', '⠉', '⠊', '⠋', '⠌', '⠍', '⠎', '⠏',
                         '⠐', '⠑', '⠒', '⠓', '⠔', '⠕', '⠖', '⠗', '⠘', '⠙', '⠚', '⠛', '⠜', '⠝', '⠞', '⠟',
                         '⠠', '⠡', '⠢', '⠣', '⠤', '⠥', '⠦', '⠧', '⠨', '⠩', '⠪', '⠫', '⠬', '⠭', '⠮', '⠯',
                         '⠰', '⠱', '⠲', '⠳', '⠴', '⠵', '⠶', '⠷', '⠸', '⠹', '⠺', '⠻', '⠼', '⠽', '⠾', '⠿']

        result = []
        for y in range(0, height * 4, 4):
            row = ''
            for x in range(0, width * 2, 2):
                pixels = []
                for dy in range(4):
                    for dx in range(2):
                        pixel_y = y + dy
                        pixel_x = x + dx
                        if pixel_y < image.shape[0] and pixel_x < image.shape[1]:
                            pixel = image[pixel_y, pixel_x]
                            pixels.append(1 if pixel < 128 else 0)
                        else:
                            pixels.append(0)
                index = (
                    pixels[0] * 1 + pixels[1] * 2 + pixels[2] * 4 + pixels[3] * 8 +
                    pixels[4] * 16 + pixels[5] * 32 + pixels[6] * 64 + pixels[7] * 128
                )
                index = min(index, len(braille_chars) - 1)
                row += braille_chars[index]
            result.append(row)
        return '\n'.join(result)

    def save_text(self, file_path, content):
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
        except Exception as e:
            raise Exception(f"Erreur lors de la sauvegarde : {str(e)}")

    def export_pdf(self, file_path, text_document, braille_text, save_type, font_name=BRAILLE_FONT_NAME, author=None, doc_name="Document"):
        """
        Export a text and/or Braille document to PDF.
        
        Args:
            file_path (str): Path to save the PDF.
            text_document: QTextDocument containing the text content.
            braille_text (str): Braille content as a string.
            save_type (str): Export type ("Texte + Braille", "Texte uniquement", "Braille uniquement").
            font_name (str): Braille font name (default: Noto Sans Braille).
            author (str): Document author (optional).
            doc_name (str): Document title (default: "Document").
        
        Raises:
            ValueError: If inputs are invalid.
            Exception: If PDF generation fails.
        """
        try:
            # Input validation
            if not file_path or not isinstance(file_path, str):
                raise ValueError("Invalid file path")
            if not text_document and save_type in ["Texte + Braille", "Texte uniquement"]:
                raise ValueError("Text document required for selected save type")
            if not braille_text and save_type in ["Texte + Braille", "Braille uniquement"]:
                raise ValueError("Braille text required for selected save type")
            if save_type not in ["Texte + Braille", "Texte uniquement", "Braille uniquement"]:
                raise ValueError("Invalid save type")

            # Step 1: Font handling
            text_font = TEXT_FONT_NAME
            braille_font = font_name

            # Vérification explicite de la présence des fichiers de police
            missing_fonts = []
            for font_key, font_path in FONT_PATHS.items():
                if not os.path.exists(font_path):
                    missing_fonts.append(f"{font_key} ({font_path})")
            if missing_fonts:
                raise Exception(f"Fichier(s) de police manquant(s) : {', '.join(missing_fonts)}.\nVeuillez vérifier les chemins dans FONT_PATHS ou installer les polices nécessaires.")

            # Register text font and its variants
            font_variants = {
                "Times New Roman": "Times New Roman",
                "Times New Roman-Bold": "Times New Roman-Bold",
                "Times New Roman-Italic": "Times New Roman-Italic",
            }
            for font_name, font_key in font_variants.items():
                if font_key not in pdfmetrics.getRegisteredFontNames():
                    font_path = FONT_PATH
                    if font_path and os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont(font_key, font_path))
                            logging.debug(f"Font '{font_key}' registered successfully from {font_path}")
                        except Exception as e:
                            logging.warning(f"Error registering font '{font_key}': {str(e)}")
                    else:
                        logging.warning(f"Font file not found for '{font_key}' at {font_path}")

            # Register Braille font
            if braille_font not in pdfmetrics.getRegisteredFontNames():
                font_path = FONT_PATHS.get(braille_font)
                if font_path and os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont(braille_font, font_path))
                        logging.debug(f"Font '{braille_font}' registered successfully from {font_path}")
                    except Exception as e:
                        logging.error(f"Error registering font '{braille_font}': {str(e)}")
                        braille_font = FALLBACK_FONT

            # Step 2: Configure PDF document with improved margins
            doc = SimpleDocTemplate(
                file_path,
                pagesize=A4,
                leftMargin=15 * mm,
                rightMargin=15 * mm,
                topMargin=20 * mm,
                bottomMargin=20 * mm,
                author=author if author else "Utilisateur non connecté",
                title=doc_name,
                subject="Conversion Texte/Braille",
                creator="Convertisseur Texte ↔ Braille"
            )

            # Step 3: Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                name="Title",
                fontName=text_font,
                fontSize=16,
                spaceAfter=20,
                spaceBefore=10,
                alignment=TA_CENTER,
                textColor='#000000'
            )
            text_style = ParagraphStyle(
                name="Text",
                fontName=text_font,
                fontSize=self.parent.base_font_size if self.parent else 12,
                spaceAfter=8,
                leading=14,
                allowWidows=1,
                allowOrphans=1,
                textColor='#333333'
            )
            braille_style = ParagraphStyle(
                name="Braille",
                fontName=braille_font,
                fontSize=self.parent.base_font_size if self.parent else 14,
                spaceAfter=8,
                leading=16,
                allowWidows=1,
                allowOrphans=1,
                textColor='#000000'
            )

            # Step 4: Initialize story and pagination parameters
            story = []
            lines_per_page = getattr(self.parent, 'lines_per_page', DEFAULT_LINES_PER_PAGE) if hasattr(self, 'parent') and self.parent else DEFAULT_LINES_PER_PAGE
            line_width = getattr(self.parent, 'line_width', DEFAULT_LINE_WIDTH) if hasattr(self, 'parent') and self.parent else DEFAULT_LINE_WIDTH
            indent_mm = getattr(self.parent, 'indent', DEFAULT_INDENT) if hasattr(self, 'parent') and self.parent else DEFAULT_INDENT
            line_spacing = getattr(self.parent, 'line_spacing', DEFAULT_LINE_SPACING) if hasattr(self, 'parent') and self.parent else DEFAULT_LINE_SPACING

            # Step 5: Add title
            # Suppression de l'affichage du titre
            # story.append(Paragraph(f"<para alignment='center'><b>{doc_name}</b></para>", title_style))
            # story.append(Spacer(1, 20))

            # Step 6: Export content based on save_type and conversion mode
            conversion_mode = getattr(self.parent, 'conversion_mode', 'text_to_braille') if hasattr(self, 'parent') and self.parent else 'text_to_braille'

            if save_type in ["Texte + Braille", "Texte uniquement"]:
                # Export text content
                block = text_document.begin()
                current_page_lines = 0
                while block.isValid():
                    if block.text().strip():
                        # Détecter l'alignement du bloc
                        align = block.blockFormat().alignment()
                        alignment = TA_LEFT
                        if int(align) & 2:
                            alignment = TA_RIGHT
                        elif int(align) & 4:
                            alignment = TA_CENTER
                        elif int(align) & 8:
                            alignment = TA_JUSTIFY

                        # Créer un style dynamique pour ce paragraphe
                        dynamic_style = ParagraphStyle(
                            name=f"Text_{alignment}",
                            parent=text_style,
                            alignment=alignment,
                            fontName=TEXT_FONT_NAME
                        )

                        # Ajouter le texte avec retour à la ligne automatique et mise en forme riche imbriquée
                        fragments = []
                        it = block.begin()
                        while not it.atEnd():
                            fragment = it.fragment()
                            if fragment.isValid():
                                frag_text = escape(fragment.text())
                                char_format = fragment.charFormat()
                                # Imbriquer les balises dans l'ordre : gras > italique > souligné
                                if char_format.fontWeight() == QFont.Bold:
                                    frag_text = f"<b>{frag_text}</b>"
                                if char_format.fontItalic():
                                    frag_text = f"<i>{frag_text}</i>"
                                if char_format.fontUnderline():
                                    frag_text = f"<u>{frag_text}</u>"
                                fragments.append(frag_text)
                            it += 1
                        html_line = "".join(fragments)
                        story.append(Paragraph(html_line, dynamic_style))
                        story.append(Spacer(1, 8))

                        # Compter les lignes et ajouter un saut de page si nécessaire
                        current_page_lines += 1
                        if current_page_lines >= lines_per_page:
                            story.append(PageBreak())
                            current_page_lines = 0

                    block = block.next()

            if save_type in ["Texte + Braille", "Braille uniquement"]:
                if save_type == "Texte + Braille":
                    story.append(PageBreak())

                # Export Braille content
                braille_lines = braille_text.split('\n')
                current_page_lines = 0
                
                for line in braille_lines:
                    if line.strip():
                        # Ajuster le style Braille en fonction du mode de conversion
                        if conversion_mode == 'braille_to_text':
                            # En mode Braille -> Texte, le texte est déjà en Braille
                            story.append(Paragraph(line, braille_style))
                        else:
                            # En mode Texte -> Braille, s'assurer que le texte est bien en Braille
                            story.append(Paragraph(line, braille_style))
                        story.append(Spacer(1, 8))
                        
                        # Compter les lignes et ajouter un saut de page si nécessaire
                        current_page_lines += 1
                        if current_page_lines >= lines_per_page:
                            story.append(PageBreak())
                            current_page_lines = 0

            # Step 7: Build the PDF
            doc.build(story)
            logging.debug(f"PDF exported successfully to {file_path}")

        except Exception as e:
            logging.error(f"Error in export_pdf: {str(e)}")
            raise Exception(f"Error exporting PDF: {str(e)}")

    def export_docx(self, file_path, text_document, braille_text, save_type, font_name=BRAILLE_FONT_NAME, doc_name="Document"):
        try:
            from docx.shared import Pt, Inches
            from PyQt5.QtGui import QFont
            doc = Document()
            from docx.oxml.ns import qn
            # Suppression de l'affichage du titre
            # doc.add_heading(doc_name, level=0)
            
            # Configurer les marges du document
            sections = doc.sections
            for section in sections:
                section.left_margin = Inches(10 / 25.4)
                section.right_margin = Inches(10 / 25.4)
                section.top_margin = Inches(10 / 25.4)
                section.bottom_margin = Inches(10 / 25.4)
            
            # Récupérer les paramètres de mise en page
            lines_per_page = self.parent.lines_per_page if self.parent else 25
            line_width = self.parent.line_width if self.parent else 80
            indent_mm = self.parent.indent if self.parent else 0
            line_spacing = self.parent.line_spacing if self.parent else 1.0
            
            if save_type in ["Texte + Braille", "Texte uniquement"]:
                block = text_document.begin()
                current_page_lines = 0
                while block.isValid():
                    if block.text().strip():
                        p = doc.add_paragraph()
                        align = block.blockFormat().alignment()
                        if int(align) & 2:
                            p.alignment = 2  # RIGHT
                        elif int(align) & 4:
                            p.alignment = 1  # CENTER
                        elif int(align) & 8:
                            p.alignment = 3  # JUSTIFY
                        else:
                            p.alignment = 0  # LEFT
                        # Appliquer le retrait et l'espacement des lignes
                        p.paragraph_format.left_indent = Inches(indent_mm / 25.4)
                        p.paragraph_format.line_spacing = line_spacing

                        # Correction : préservation des espaces entre fragments stylisés
                        it = block.begin()
                        last_end = 0
                        block_text = block.text()
                        while not it.atEnd():
                            fragment = it.fragment()
                            if fragment.isValid():
                                char_format = fragment.charFormat()
                                frag_text = fragment.text()
                                start = fragment.position() - block.position()
                                # Ajouter les espaces ou texte intermédiaire non stylisé
                                if start > last_end:
                                    p.add_run(block_text[last_end:start])
                                run = p.add_run(frag_text)
                                run.font.name = self.parent.current_font if self.parent else font_name
                                fragment_font_size = char_format.fontPointSize()
                                if fragment_font_size > 0:
                                    run.font.size = Pt(fragment_font_size)
                                else:
                                    run.font.size = Pt(self.parent.base_font_size if self.parent else 12)
                                run.bold = char_format.fontWeight() == QFont.Bold
                                run.italic = char_format.fontItalic()
                                run.underline = char_format.fontUnderline()
                                last_end = start + len(frag_text)
                            it += 1
                        # Ajouter le reste du texte s'il y en a
                        if last_end < len(block_text):
                            p.add_run(block_text[last_end:])
                        # Compter les lignes et ajouter un saut de page si nécessaire
                        wrapped_text = self._wrap_text(block.text(), line_width)
                        current_page_lines += len(wrapped_text.split('\n'))
                        if current_page_lines >= lines_per_page:
                            p.add_run().add_break(WD_BREAK.PAGE)
                            current_page_lines = 0
                    block = block.next()

            if save_type in ["Texte + Braille", "Braille uniquement"]:
                if save_type == "Texte + Braille":
                    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                
                braille_lines = braille_text.split('\n')
                current_page_lines = 0
                
                for line in braille_lines:
                    if line.strip():
                        p = doc.add_paragraph()
                        formatted_line = self._wrap_text(line, line_width)
                        run = p.add_run(formatted_line)
                        run.font.name = font_name
                        p.paragraph_format.left_indent = Inches(indent_mm / 25.4)
                        p.paragraph_format.line_spacing = line_spacing
                        
                        # Compter les lignes et ajouter un saut de page si nécessaire
                        current_page_lines += 1
                        if current_page_lines >= lines_per_page:
                            p.add_run().add_break(WD_BREAK.PAGE)
                            current_page_lines = 0

            doc.save(file_path)
            print(f"Document DOCX exporté avec succès : {file_path}")
        except Exception as e:
            raise Exception(f"Erreur lors de l'exportation en DOCX : {str(e)}")

    def export_to_gcode(self, file_path):
        if self.last_gcode is None:
            print("Aucun G-code disponible pour l'exportation.")
            return False
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(self.last_gcode)
            print(f"G-code exporté avec succès : {file_path}")
            return True
        except Exception as e:
            print(f"Erreur lors de l'exportation du G-code : {str(e)}")
            return False

    def print_content(self, printer, text_content, braille_content):
        try:
            painter = QPainter()
            if not painter.begin(printer):
                print("Erreur : Impossible d'initialiser l'impression.")
                return False

            font = QFont(BRAILLE_FONT_NAME, 12)
            painter.setFont(font)
            lines_per_page = self.parent.lines_per_page if self.parent else 25
            line_width = self.parent.line_width if self.parent else 80
            indent_pixels = (self.parent.indent * 2.83) if self.parent else 0
            line_spacing = self.parent.line_spacing if self.parent else 1.0

            y = 50
            page_height = printer.pageRect().height() - 100
            line_height = painter.fontMetrics().height() * line_spacing
            max_lines_per_page = min(lines_per_page, int(page_height / line_height))

            if text_content.strip():
                text_lines = text_content.split('\n')
                current_page_lines = []
                line_count = 0

                for line in text_lines:
                    if line_count >= max_lines_per_page:
                        printer.newPage()
                        y = 50
                        for saved_line in current_page_lines:
                            wrapped_line = self._wrap_text(saved_line, line_width)
                            painter.drawText(50 + indent_pixels, y, wrapped_line)
                            y += line_height
                        current_page_lines = []
                        line_count = 0
                    if line.strip():
                        current_page_lines.append(line)
                        line_count += 1

                if current_page_lines:
                    printer.newPage()
                    y = 50
                    for saved_line in current_page_lines:
                        wrapped_line = self._wrap_text(saved_line, line_width)
                        painter.drawText(50 + indent_pixels, y, wrapped_line)
                        y += line_height

            if braille_content.strip():
                if text_content.strip():
                    printer.newPage()
                    y = 50
                    painter.drawText(50, y, "=== Section Braille ===")
                    y += line_height * 2

                braille_lines = braille_content.split('\n')
                current_page_lines = []
                line_count = 0

                for line in braille_lines:
                    if line_count >= max_lines_per_page:
                        printer.newPage()
                        y = 50
                        for saved_line in current_page_lines:
                            wrapped_line = self._wrap_text(saved_line, line_width)
                            painter.drawText(50 + indent_pixels, y, wrapped_line)
                            y += line_height
                        current_page_lines = []
                        line_count = 0
                    if line.strip():
                        current_page_lines.append(line)
                        line_count += 1

                if current_page_lines:
                    printer.newPage()
                    y = 50
                    for saved_line in current_page_lines:
                        wrapped_line = self._wrap_text(saved_line, line_width)
                        painter.drawText(50 + indent_pixels, y, wrapped_line)
                        y += line_height

            painter.end()
            print("Contenu imprimé avec succès.")
            return True

        except Exception as e:
            print(f"Erreur lors de l'impression : {str(e)}")
            return False

    def _wrap_text(self, text, max_width):
        logging.debug(f"_wrap_text called with text='{text[:50]}...', max_width={max_width}")
        if not text or max_width < 1:
            return text

        lines = []
        current_line = ""
        segments = re.split(r'(\s+|[^\s]+)', text)
        segments = [s for s in segments if s]

        for segment in segments:
            if segment.isspace():
                if len(current_line) + len(segment) <= max_width:
                    current_line += segment
                else:
                    if current_line:
                        lines.append(current_line.rstrip())
                    current_line = segment.lstrip()
                continue

            if len(current_line) + len(segment) <= max_width:
                current_line += segment
            else:
                if current_line:
                    lines.append(current_line.rstrip())
                    current_line = segment
                else:
                    while len(segment) > max_width:
                        lines.append(segment[:max_width])
                        segment = segment[max_width:]
                    current_line = segment

        if current_line:
            lines.append(current_line.rstrip())

        result = "\n".join(lines)
        logging.debug(f"Wrapped text: {result[:100]}...")
        return result

    def convert_to_gcode(self, text):
        if not text.strip():
            return "; Aucun contenu à convertir en G-code\n"

        gcode_lines = []
        gcode_lines.append("; G-code généré à partir du texte Braille")
        gcode_lines.append("G21 ; Utiliser des unités en millimètres")
        gcode_lines.append("G90 ; Utiliser un positionnement absolu")
        gcode_lines.append("G0 Z5.0 ; Lever l'outil")
        gcode_lines.append("G0 X0 Y0 ; Aller à la position initiale")

        x, y = 0, 0
        for char in text:
            if char == '\n':
                y -= 5
                x = 0
            else:
                gcode_lines.append(f"G0 X{x} Y{y} ; Position pour caractère")
                gcode_lines.append("G1 Z0 ; Abaisser l'outil")
                gcode_lines.append("G1 Z5 ; Lever l'outil")
                x += 2

        gcode_lines.append("G0 X0 Y0 ; Retour à l'origine")
        gcode_lines.append("; Fin du G-code")
        return "\n".join(gcode_lines)